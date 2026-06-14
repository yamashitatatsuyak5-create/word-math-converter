import io
import re
import streamlit as st
from docx import Document
import math2docx

def convert_math_in_paragraph(paragraph):
    """
    段落内の $$...$$ (ブロック) と $...$ (文中) を解析し、
    math2docxを使って本物のWord数式オブジェクトに完全変換する
    """
    text = paragraph.text
    if not text:
        return

    # $$...$$ と $...$ を同時に切り分ける正規表現
    pattern = r'(\$\$.*?\$\$|\$.*?\$)'
    tokens = re.split(pattern, text)
    
    # 数式が含まれていない場合はスキップ
    if len(tokens) == 1:
        return

    # 段落のテキストを一度空にして、パーツごとに再構築する
    paragraph.text = ""
    
    for token in tokens:
        if not token:
            continue
        
        if token.startswith("$$") and token.endswith("$$"):
            # 1. 独立行の数式（$$...$$）
            math_content = token[2:-2]
            try:
                # math2docxがLaTeXをWord数式(OMML)に翻訳して段落に追加
                math2docx.add_math(paragraph, math_content)
            except Exception:
                # 万が一翻訳できない複雑な数式の場合はそのまま文字で返す
                paragraph.add_run(f"$$ {math_content} $$")
                
        elif token.startswith("$") and token.endswith("$"):
            # 2. 文中の数式（$...$）
            math_content = token[1:-1]
            try:
                math2docx.add_math(paragraph, math_content)
            except Exception:
                paragraph.add_run(f"$ {math_content} $")
                
        else:
            # 3. 通常のテキスト
            paragraph.add_run(token)

def process_docx(file_bytes):
    """
    Wordファイルを読み込み、すべての段落の数式を変換する
    """
    doc = Document(io.BytesIO(file_bytes))
    
    # 通常の段落を処理
    for paragraph in doc.paragraphs:
        convert_math_in_paragraph(paragraph)
        
    # 表（テーブル）の中の段落も処理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    convert_math_in_paragraph(paragraph)
                    
    # 変換したファイルをメモリ上に保存
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# --- Streamlit UI の構築 ---
st.title("Word数式 完全自動変換アプリ 📝✨")
st.write("分数の `\\frac{a}{b}` なども完璧にWordの本物の数式に変換します！")

uploaded_file = st.file_uploader("Wordファイル (.docx) をアップロードしてください", type=["docx"])

if uploaded_file is not None:
    st.success("ファイルを読み込みました！変換を開始します...")
    
    # ファイル処理
    file_bytes = uploaded_file.read()
    processed_docx = process_docx(file_bytes)
    
    # ダウンロードボタン
    st.download_button(
        label="変換済みWordファイルをダウンロード",
        data=processed_docx,
        file_name=f"converted_{uploaded_file.name}",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )